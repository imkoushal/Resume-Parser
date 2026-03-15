"""
text_extractor.py - Text Extraction module.

This module handles raw text extraction from resume files in PDF and DOCX
formats, returning clean, normalized text ready for further processing.

Supported formats:
    - PDF  (.pdf)  — via pdfplumber
    - Word (.docx) — via python-docx
    - Text (.txt)  — built-in

Usage:
    # Convenience function (recommended)
    text = extract_text("resume.pdf")

    # Class-based usage
    extractor = TextExtractor()
    text = extractor.extract("resume.docx")
"""

import os
import re

import pdfplumber
from docx import Document


# ---------------------------------------------------------------------------
# Convenience function (top-level API)
# ---------------------------------------------------------------------------

def extract_text(file) -> str:
    """
    Extract and return clean text from a resume file.

    Accepts either a file path string or a file-like object.  The format is
    determined by the file extension (for paths) or the ``name`` attribute
    (for file-like objects).

    Args:
        file (str | os.PathLike | file-like): Path to the resume file, or an
            open binary file-like object whose ``.name`` attribute ends with
            a supported extension.

    Returns:
        str: Clean, whitespace-normalised text extracted from the resume.

    Raises:
        FileNotFoundError: If *file* is a path that does not exist.
        ValueError: If the file extension is not supported.
    """
    extractor = TextExtractor()
    return extractor.extract(file)


# ---------------------------------------------------------------------------
# TextExtractor class
# ---------------------------------------------------------------------------

class TextExtractor:
    """
    Extracts raw text from resume files of various formats.

    Supported formats:
        - PDF  (.pdf)
        - Word (.docx)
        - Text (.txt)
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------

    def extract(self, file) -> str:
        """
        Extract raw text from a file, auto-detecting the format.

        Args:
            file (str | os.PathLike | file-like): Path to the resume file or
                an open binary file-like object with a ``.name`` attribute.

        Returns:
            str: Cleaned extracted text.

        Raises:
            FileNotFoundError: If a path is given and the file does not exist.
            ValueError: If the file extension is not supported.
        """
        # Resolve the file name / extension
        if hasattr(file, "read"):
            # File-like object (e.g. from open() or an upload stream)
            name = getattr(file, "name", "")
        else:
            name = str(file)
            if not os.path.isfile(name):
                raise FileNotFoundError(f"File not found: {name}")

        ext = os.path.splitext(name)[1].lower()

        if ext == ".pdf":
            raw = self.extract_from_pdf(file)
        elif ext == ".docx":
            raw = self.extract_from_docx(file)
        elif ext == ".txt":
            raw = self.extract_from_txt(file)
        else:
            raise ValueError(
                f"Unsupported file format '{ext}'. "
                f"Supported formats: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        return self.clean_text(raw)

    # ------------------------------------------------------------------
    # Format-specific extractors
    # ------------------------------------------------------------------

    def extract_from_pdf(self, file) -> str:
        """
        Extract text from a PDF file using *pdfplumber*.

        All pages are concatenated with a newline separator.  Tables are
        extracted as tab-separated rows so that structured data (e.g. skill
        tables) is not lost.

        Args:
            file (str | os.PathLike | file-like): Path or binary file-like
                object pointing to a PDF.

        Returns:
            str: Raw text extracted from all pages.
        """
        pages_text: list[str] = []

        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                # Extract plain text from the page
                page_text = page.extract_text() or ""

                # Also extract any tables and append as plain text rows
                tables = page.extract_tables() or []
                table_lines: list[str] = []
                for table in tables:
                    for row in table:
                        # Filter None cells and join with tab
                        row_text = "\t".join(cell or "" for cell in row)
                        if row_text.strip():
                            table_lines.append(row_text)

                combined = page_text
                if table_lines:
                    combined = combined + "\n" + "\n".join(table_lines)

                pages_text.append(combined)

        return "\n".join(pages_text)

    def extract_from_docx(self, file) -> str:
        """
        Extract text from a DOCX file using *python-docx*.

        Text is gathered from:
        - All body paragraphs (preserving order)
        - All table cells (row by row, cell by cell)

        Args:
            file (str | os.PathLike | file-like): Path or binary file-like
                object pointing to a DOCX document.

        Returns:
            str: Raw text extracted from the document.
        """
        doc = Document(file)
        lines: list[str] = []

        # Body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                row_text = "\t".join(cell for cell in row_cells if cell)
                if row_text.strip():
                    lines.append(row_text)

        return "\n".join(lines)

    def extract_from_txt(self, file) -> str:
        """
        Extract text from a plain text file.

        Args:
            file (str | os.PathLike | file-like): Path or text/binary
                file-like object pointing to a TXT file.

        Returns:
            str: Content of the text file.
        """
        if hasattr(file, "read"):
            raw = file.read()
            # Handle binary mode
            if isinstance(raw, bytes):
                raw = raw.decode("utf-8", errors="replace")
            return raw

        with open(file, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    # ------------------------------------------------------------------
    # Text cleaning
    # ------------------------------------------------------------------

    def clean_text(self, text: str) -> str:
        """
        Clean and normalise extracted text.

        Steps applied:
        1. Replace Windows-style line endings with Unix style.
        2. Remove non-printable / control characters (except newlines/tabs).
        3. Collapse runs of blank lines into a single blank line.
        4. Strip leading/trailing whitespace from each line.
        5. Strip leading/trailing whitespace from the whole text.

        Args:
            text (str): Raw extracted text.

        Returns:
            str: Cleaned and normalised text.
        """
        if not text:
            return ""

        # 1. Normalise line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # 2. Remove non-printable control characters (keep \n and \t)
        text = re.sub(r"[^\S\n\t ]+", " ", text)          # collapse odd whitespace
        text = re.sub(r"[^\x09\x0A\x20-\x7E\u00A0-\uFFFF]", "", text)

        # 3. Strip each line
        lines = [line.strip() for line in text.split("\n")]

        # 4. Collapse multiple consecutive blank lines into one
        cleaned_lines: list[str] = []
        prev_blank = False
        for line in lines:
            is_blank = line == ""
            if is_blank and prev_blank:
                continue
            cleaned_lines.append(line)
            prev_blank = is_blank

        return "\n".join(cleaned_lines).strip()
